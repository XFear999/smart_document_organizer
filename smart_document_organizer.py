#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Smart Document Organizer
========================

A safe, accurate, rule-based document organizer for Windows (also works on
macOS / Linux). It scans a folder (and subfolders), extracts text from many
document types (using OCR for scanned PDFs and images), classifies each file
into a category using *strong* rule-based scoring (not naive keyword matching),
detects exact duplicates by SHA-256, and then either reports what it *would*
do (dry-run, the default) or copies / moves files into a clean, organized
folder tree.

Design goals
------------
* SAFE BY DEFAULT
    - Dry-run is the default. Nothing is touched unless ``--apply`` is passed.
    - Copy mode is the recommended safe operation; move mode is available but
      must be explicitly requested.
    - Files are NEVER deleted automatically.
    - Existing files are NEVER overwritten. Collisions get ``(2)``, ``(3)`` ...
* ACCURATE
    - Each category has must-have signals, strong/weak keywords, negative
      keywords, and confidence scoring. Weak matches are sent to
      ``_Needs_Review`` (or just logged) instead of being forced into a folder.
* AUDITABLE
    - Every decision is written to a CSV log and an SQLite database with full
      detail (score breakdown, detected fields, duplicate info, text preview).
* SELF-IMPROVING
    - You can correct the ``corrected_category`` column in the CSV and feed it
      back with ``--learn-from``. Corrections are keyed by SHA-256 so the same
      file (by content) is classified correctly next time.

This module intentionally degrades gracefully: optional third-party libraries
(pdfplumber, python-docx, openpyxl, pytesseract, ...) are imported lazily and
their absence only disables the corresponding feature; the program keeps
running for everything else.

Author: generated for Bassel Khair
License: MIT
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import io
import logging
import os
import re
import sqlite3
import sys
import shutil
import zipfile
from dataclasses import dataclass, field, asdict
from datetime import datetime, date
from pathlib import Path
from typing import Callable, Dict, List, Optional, Tuple

# --------------------------------------------------------------------------- #
#  Version / constants
# --------------------------------------------------------------------------- #

APP_NAME = "Smart Document Organizer"
APP_VERSION = "1.1.0"

# Files that live in the organizer's own bookkeeping. We never move these.
LOG_CSV_NAME = "organizer_log.csv"
DB_NAME = "organizer.db"

# Name of the top-level output tree.
DEFAULT_OUTPUT_DIRNAME = "Organized_Documents"

# Folders we should never descend into while scanning (system / noise).
SKIP_DIR_NAMES = {
    "$recycle.bin",
    "system volume information",
    ".git",
    ".svn",
    "__pycache__",
    "node_modules",
    "$windows.~bt",
    "$windows.~ws",
    "windows",
    "program files",
    "program files (x86)",
    "appdata",
}

# Extensions we know how to handle. Anything else is logged but left alone
# (classified as "Needs Review" with a note).
SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".doc",   # best-effort (may need textract/antiword; we try anyway)
    ".xlsx",
    ".xls",
    ".csv",
    ".txt",
    ".jpg",
    ".jpeg",
    ".png",
    ".tif",
    ".tiff",
    ".webp",
    ".eml",
    ".msg",
    ".zip",
}

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".webp"}

# Amount of extracted text we keep as a preview in logs.
TEXT_PREVIEW_LEN = 400
# Max characters of text we classify on (keeps things fast for huge files).
MAX_TEXT_FOR_CLASSIFY = 20000

logger = logging.getLogger("organizer")


# --------------------------------------------------------------------------- #
#  Categories, folder layout, and the classification rule-set
# --------------------------------------------------------------------------- #

# Special buckets that are not "classified" in the normal sense.
CAT_NEEDS_REVIEW = "Needs Review"
CAT_DUPLICATE = "Exact Duplicates"

# Maps a category -> its relative folder path inside the output tree.
# (The final /Bank/Year sub-levels are appended dynamically at move time.)
CATEGORY_FOLDER: Dict[str, str] = {
    "Bank Statements": r"Financial\Bank Statements",
    "Credit Card Statements": r"Financial\Credit Card Statements",
    "Invoices": r"Financial\Invoices",
    "Receipts": r"Financial\Receipts",
    "IRS Notices": r"Tax\IRS Notices",
    "Tax Forms": r"Tax\Tax Forms",
    "Contracts": r"Legal\Contracts",
    "Drafts": r"Legal\Drafts",
    "Legal Filings": r"Legal\Court Filings",
    "Insurance": r"Insurance",
    "Medical": r"Medical",
    "Identity Documents": r"Personal\Identity Documents",
    "Immigration and Travel": r"Personal\Immigration and Travel",
    "Real Estate": r"Real Estate",
    "Business Documents": r"Business\Business Documents",
    "Payroll and HR": r"Business\Payroll and HR",
    "Vehicle Documents": r"Vehicles",
    "Utilities": r"Utilities and Bills",
    "Personal": r"Personal",
    CAT_NEEDS_REVIEW: r"_Needs_Review",
    CAT_DUPLICATE: r"_Duplicates",
}


@dataclass
class Rule:
    """A scoring rule-set for one category.

    Scoring model (transparent and tunable):

      * ``must_have``   -> at least ``min_must_have`` of these phrase-groups
                           must be present or the category is disqualified
                           (score forced to 0). This is what stops weak/random
                           matches from being forced into a folder.
      * ``strong``      -> +STRONG_WEIGHT each (these are high-signal phrases)
      * ``weak``        -> +WEAK_WEIGHT each (supporting evidence)
      * ``negative``    -> -NEG_WEIGHT each (evidence *against* this category)

    Each entry may itself be a tuple of synonyms; matching any synonym counts
    the group once (so "amount due"/"balance due" don't double count).
    """

    name: str
    must_have: List = field(default_factory=list)
    strong: List = field(default_factory=list)
    weak: List = field(default_factory=list)
    negative: List = field(default_factory=list)
    min_must_have: int = 1
    # Per-category base to nudge priority when two categories tie.
    base: float = 0.0


# Weights for the scoring model. Tuned so that a document with the required
# must-have signals plus a couple of strong phrases comfortably clears the
# default confidence threshold, while a document matching only stray weak
# keywords does not.
STRONG_WEIGHT = 2.0
WEAK_WEIGHT = 0.6
NEG_WEIGHT = 1.5
MUST_HAVE_WEIGHT = 2.5  # each satisfied must-have group also adds score

# The maximum "raw" score we normalize against when computing confidence.
# Confidence = min(1.0, raw_score / CONFIDENCE_FULL_SCALE).
CONFIDENCE_FULL_SCALE = 10.0


def _build_rules() -> Dict[str, Rule]:
    """Construct the full rule-set. Kept in a function for readability."""
    rules: Dict[str, Rule] = {}

    rules["Bank Statements"] = Rule(
        name="Bank Statements",
        must_have=[
            ("statement period", "statement date", "for the period"),
            ("beginning balance", "opening balance", "previous balance"),
            ("ending balance", "closing balance", "new balance"),
        ],
        min_must_have=2,
        strong=[
            "account summary",
            ("deposits and credits", "total deposits", "deposits"),
            ("withdrawals and debits", "total withdrawals", "withdrawals"),
            ("checking account", "savings account"),
            "routing number",
            ("account number", "acct number", "acct #"),
            "available balance",
            "direct deposit",
        ],
        weak=["transaction", "balance", "statement", "interest earned",
              "service fee", "minimum balance"],
        negative=["invoice number", "amount due", "credit card statement",
                  "minimum payment due", "explanation of benefits"],
        base=0.3,
    )

    rules["Credit Card Statements"] = Rule(
        name="Credit Card Statements",
        must_have=[
            ("minimum payment due", "minimum payment"),
            ("payment due date", "due date"),
            ("new balance", "statement balance", "closing balance"),
        ],
        min_must_have=2,
        strong=[
            "credit limit",
            "available credit",
            ("purchases", "cash advances"),
            "annual percentage rate",
            ("apr", "finance charge"),
            "credit card statement",
            "rewards",
            "previous balance",
        ],
        weak=["statement", "transaction", "interest", "late payment",
              "card ending"],
        negative=["routing number", "checking account", "savings account",
                  "invoice number"],
        base=0.3,
    )

    rules["IRS Notices"] = Rule(
        name="IRS Notices",
        must_have=[
            ("internal revenue service", "department of the treasury",
             "department of treasury"),
        ],
        min_must_have=1,
        strong=[
            re.compile(r"\bnotice\s+cp\d{2,4}\b", re.I),
            re.compile(r"\bletter\s+\d{3,4}\b", re.I),
            "form 4564",
            ("information document request", "idr"),
            ("examination", "audit"),
            "we changed your",
            "amount you owe",
            "taxpayer identification",
        ],
        weak=["irs", "tax year", "notice", "response form", "penalty",
              "interest charged"],
        negative=["invoice number", "bill to", "purchase order"],
        base=0.4,
    )

    rules["Tax Forms"] = Rule(
        name="Tax Forms",
        must_have=[
            (
                re.compile(r"\bform\s+1040\b", re.I),
                re.compile(r"\bform\s+1120\b", re.I),
                re.compile(r"\bform\s+1065\b", re.I),
                re.compile(r"\bform\s+990\b", re.I),
                re.compile(r"\bw-?2\b", re.I),
                re.compile(r"\bw-?9\b", re.I),
                re.compile(r"\b1099(-[a-z]{1,4})?\b", re.I),
                re.compile(r"\bschedule\s+k-?1\b", re.I),
                re.compile(r"\b1040\b", re.I),
            ),
        ],
        min_must_have=1,
        strong=[
            "u.s. individual income tax return",
            "wage and tax statement",
            "request for taxpayer identification",
            "adjusted gross income",
            "taxable income",
            "employer identification number",
            "withholding",
        ],
        weak=["tax", "irs", "income", "deduction", "filing status",
              "gross income"],
        negative=["invoice number", "bill to", "statement period",
                  "notice cp"],
        base=0.35,
    )

    rules["Invoices"] = Rule(
        name="Invoices",
        must_have=[
            (re.compile(r"invoice\s*(no\.?|number|#)", re.I), "invoice #"),
            ("amount due", "balance due", "total due", "please pay"),
        ],
        min_must_have=2,
        strong=[
            ("bill to", "billed to"),
            "payment terms",
            ("subtotal", "sub total"),
            ("total due", "grand total", "invoice total"),
            "net 30",
            "purchase order",
            "remit to",
            "tax id",
        ],
        weak=["invoice", "quantity", "unit price", "description", "due date",
              "amount"],
        negative=["statement period", "minimum payment due", "receipt",
                  "beginning balance"],
        base=0.2,
    )

    rules["Receipts"] = Rule(
        name="Receipts",
        must_have=[
            ("receipt", "sales receipt", "payment received", "paid",
             "thank you for your purchase", "order confirmation"),
        ],
        min_must_have=1,
        strong=[
            ("total paid", "amount paid", "you paid"),
            ("subtotal", "tax", "change due"),
            ("card ending", "visa", "mastercard", "amex", "cash tendered"),
            "transaction id",
            "merchant",
            "order number",
        ],
        weak=["receipt", "total", "qty", "item", "store", "purchase"],
        negative=["invoice number", "amount due", "statement period",
                  "minimum payment"],
        base=0.1,
    )

    rules["Contracts"] = Rule(
        name="Contracts",
        must_have=[
            ("agreement", "contract", "this agreement"),
            ("whereas", "effective date", "the parties", "by and between",
             "signature", "signed by", "in witness whereof"),
        ],
        min_must_have=2,
        strong=[
            "effective date",
            ("by and between", "between the parties"),
            "whereas",
            ("in witness whereof", "signature of"),
            ("terms and conditions", "governing law"),
            "hereby agree",
            "indemnification",
            ("licensing agreement", "service agreement",
             "non-disclosure agreement", "nda"),
        ],
        weak=["party", "obligation", "term", "clause", "shall", "hereof",
              "consideration"],
        negative=["invoice number", "amount due", "docket", "plaintiff",
                  "statement period"],
        base=0.15,
    )

    rules["Drafts"] = Rule(
        name="Drafts",
        must_have=[
            ("draft", "do not distribute", "working copy", "for discussion",
             "not for execution", "redline"),
        ],
        min_must_have=1,
        strong=["draft", "confidential draft", "for review only",
                "subject to change", "preliminary"],
        weak=["revision", "version", "comment", "tbd", "placeholder"],
        negative=["executed", "fully signed", "final version"],
        base=0.05,
    )

    rules["Legal Filings"] = Rule(
        name="Legal Filings",
        must_have=[
            ("court", "in the matter of", "united states district court",
             "superior court", "district court"),
            ("plaintiff", "defendant", "docket", "case no", "motion",
             "complaint", "summons", "petitioner", "respondent"),
        ],
        min_must_have=2,
        strong=[
            ("docket no", "case no", "civil action no"),
            ("plaintiff", "defendant"),
            ("motion", "order", "summons", "complaint", "subpoena"),
            "certification",
            "hereby ordered",
            "comes now",
            "notice of hearing",
        ],
        weak=["court", "counsel", "attorney", "filed", "judge", "jurisdiction"],
        negative=["invoice number", "amount due", "statement period"],
        base=0.25,
    )

    rules["Insurance"] = Rule(
        name="Insurance",
        must_have=[
            ("policy number", "policy no", "insurance policy",
             "certificate of insurance", "declarations page"),
        ],
        min_must_have=1,
        strong=[
            ("premium", "annual premium"),
            ("coverage", "deductible", "coverage limit"),
            ("insured", "policyholder", "named insured"),
            ("claim number", "claims"),
            "underwriting",
            "explanation of benefits",
            "beneficiary",
        ],
        weak=["insurance", "policy", "liability", "endorsement", "rider"],
        negative=["invoice number", "docket", "statement period"],
        base=0.2,
    )

    rules["Medical"] = Rule(
        name="Medical",
        must_have=[
            ("patient", "diagnosis", "prescription", "medical record",
             "explanation of benefits", "lab results", "physician",
             "date of service"),
        ],
        min_must_have=1,
        strong=[
            ("patient name", "date of birth", "mrn"),
            ("diagnosis", "icd-10", "cpt"),
            ("prescription", "rx", "dosage"),
            ("provider", "physician", "clinic", "hospital"),
            "explanation of benefits",
            "date of service",
            "immunization",
        ],
        weak=["medical", "health", "treatment", "symptom", "referral",
              "copay"],
        negative=["invoice number", "docket", "statement period",
                  "policy number"],
        base=0.2,
    )

    rules["Identity Documents"] = Rule(
        name="Identity Documents",
        must_have=[
            ("passport", "driver's license", "driver license",
             "identification card", "social security", "birth certificate",
             "national id", "state id"),
        ],
        min_must_have=1,
        strong=[
            ("passport no", "passport number"),
            ("date of birth", "dob"),
            ("license number", "dln", "id number"),
            ("social security number", "ssn"),
            "place of birth",
            "nationality",
            "sex",
        ],
        weak=["identity", "issued", "expiration", "authority", "surname",
              "given name"],
        negative=["invoice number", "docket", "statement period"],
        base=0.2,
    )

    rules["Immigration and Travel"] = Rule(
        name="Immigration and Travel",
        must_have=[
            ("visa", "i-94", "i-797", "green card", "permanent resident",
             "uscis", "boarding pass", "itinerary", "customs",
             "employment authorization", "ds-160"),
        ],
        min_must_have=1,
        strong=[
            ("uscis", "department of homeland security"),
            ("i-94", "i-797", "i-140", "i-485", "ds-160"),
            ("visa", "entry visa", "nonimmigrant visa"),
            ("boarding pass", "flight", "itinerary", "pnr"),
            "permanent resident card",
            "employment authorization document",
            "port of entry",
        ],
        weak=["travel", "passport", "airline", "departure", "arrival",
              "immigration"],
        negative=["invoice number", "statement period"],
        base=0.2,
    )

    rules["Real Estate"] = Rule(
        name="Real Estate",
        must_have=[
            ("lease", "deed", "mortgage", "closing disclosure",
             "purchase agreement", "title insurance", "property",
             "landlord", "tenant", "escrow"),
        ],
        min_must_have=1,
        strong=[
            ("purchase agreement", "closing disclosure", "settlement statement"),
            ("mortgage", "promissory note", "deed of trust"),
            ("landlord", "tenant", "lessor", "lessee"),
            ("escrow", "title", "parcel"),
            "property address",
            "square feet",
            "homeowners association",
        ],
        weak=["property", "rent", "real estate", "premises", "county",
              "appraisal"],
        negative=["invoice number", "docket", "policy number"],
        base=0.2,
    )

    rules["Business Documents"] = Rule(
        name="Business Documents",
        must_have=[
            ("business plan", "proposal", "memorandum", "articles of",
             "operating agreement", "meeting minutes", "purchase order",
             "statement of work", "sow", "non-disclosure",
             "financial statement", "balance sheet", "profit and loss"),
        ],
        min_must_have=1,
        strong=[
            ("articles of incorporation", "articles of organization",
             "operating agreement"),
            ("statement of work", "scope of work", "sow"),
            ("balance sheet", "profit and loss", "income statement",
             "cash flow"),
            ("meeting minutes", "board resolution"),
            "purchase order",
            "employer identification number",
            "business plan",
        ],
        weak=["company", "business", "corporate", "vendor", "quarterly",
              "fiscal"],
        negative=["patient", "docket", "passport"],
        base=0.05,
    )

    rules["Payroll and HR"] = Rule(
        name="Payroll and HR",
        must_have=[
            ("pay stub", "paystub", "payroll", "earnings statement",
             "offer letter", "employment agreement", "w-2", "gross pay",
             "net pay", "direct deposit advice", "benefits enrollment"),
        ],
        min_must_have=1,
        strong=[
            ("gross pay", "net pay", "ytd", "year to date"),
            ("pay period", "pay date"),
            ("employee id", "employee number"),
            ("withholding", "deductions", "fica"),
            ("offer letter", "employment agreement", "onboarding"),
            "benefits enrollment",
            "time sheet",
        ],
        weak=["employee", "employer", "salary", "wage", "hours", "hr"],
        negative=["invoice number", "docket", "patient"],
        base=0.15,
    )

    rules["Vehicle Documents"] = Rule(
        name="Vehicle Documents",
        must_have=[
            ("vehicle", "vin", "registration", "title certificate",
             "odometer", "license plate", "auto insurance", "make and model",
             "bill of sale"),
        ],
        min_must_have=1,
        strong=[
            re.compile(r"\bvin\b", re.I),
            ("make", "model", "year", "odometer"),
            ("registration", "certificate of title"),
            ("license plate", "plate number"),
            "bill of sale",
            "dmv",
            "vehicle identification number",
        ],
        weak=["vehicle", "car", "truck", "auto", "mileage", "dealer"],
        negative=["invoice number", "docket", "patient", "policy number"],
        base=0.15,
    )

    rules["Utilities"] = Rule(
        name="Utilities",
        must_have=[
            ("utility", "electric bill", "water bill", "gas bill",
             "internet bill", "cable", "kwh", "meter reading",
             "service address", "account number"),
            ("amount due", "total due", "current charges", "past due",
             "auto pay", "billing period"),
        ],
        min_must_have=2,
        strong=[
            ("kwh", "meter reading", "usage"),
            ("service address", "service period"),
            ("current charges", "previous balance", "amount due"),
            ("electric", "water", "gas", "internet", "cable", "phone"),
            "utility",
            "billing period",
        ],
        weak=["bill", "provider", "account", "due date", "payment"],
        negative=["invoice number", "docket", "statement period",
                  "minimum payment due"],
        base=0.1,
    )

    rules["Personal"] = Rule(
        name="Personal",
        must_have=[
            ("dear", "letter", "resume", "cv", "curriculum vitae",
             "personal", "note", "diary", "certificate", "diploma",
             "transcript"),
        ],
        min_must_have=1,
        strong=["resume", "curriculum vitae", "diploma", "certificate of",
                "transcript", "cover letter"],
        weak=["personal", "family", "note", "letter", "photo", "school"],
        negative=["invoice number", "docket", "statement period"],
        base=0.0,
    )

    return rules


RULES: Dict[str, Rule] = _build_rules()


# --------------------------------------------------------------------------- #
#  Party / bank / date detection helpers
# --------------------------------------------------------------------------- #

# A small curated list of well-known banks / card issuers so we can produce a
# clean folder name (e.g. "Chase") and a good rename token. Extend freely.
KNOWN_BANKS = [
    "Chase", "JPMorgan", "Bank of America", "Wells Fargo", "Citibank",
    "Citi", "Capital One", "American Express", "Amex", "US Bank",
    "U.S. Bank", "PNC", "TD Bank", "Truist", "Discover", "Barclays",
    "HSBC", "Goldman Sachs", "Marcus", "Ally", "Charles Schwab", "Fidelity",
    "SoFi", "Synchrony", "Navy Federal", "USAA", "Regions", "Fifth Third",
    "KeyBank", "Santander", "Citizens Bank", "M&T Bank", "BB&T", "Chime",
    "Venmo", "PayPal", "Robinhood", "Comerica", "First Republic",
]

# Common date patterns to search for a document date. Ordered by specificity.
_MONTHS = (
    "january|february|march|april|may|june|july|august|september|october|"
    "november|december|jan|feb|mar|apr|jun|jul|aug|sep|sept|oct|nov|dec"
)

DATE_PATTERNS = [
    # 2025-03-31  /  2025/03/31
    (re.compile(r"\b(20\d{2})[-/\.](\d{1,2})[-/\.](\d{1,2})\b"), "ymd"),
    # 03/31/2025  /  03-31-2025  (US month-first)
    (re.compile(r"\b(\d{1,2})[-/\.](\d{1,2})[-/\.](20\d{2})\b"), "mdy"),
    # March 31, 2025  /  Mar 31 2025
    (re.compile(r"\b(" + _MONTHS + r")\.?\s+(\d{1,2}),?\s+(20\d{2})\b", re.I),
     "month_name"),
    # 31 March 2025
    (re.compile(r"\b(\d{1,2})\s+(" + _MONTHS + r")\.?\s+(20\d{2})\b", re.I),
     "day_month_name"),
]

_MONTH_INDEX = {
    "january": 1, "jan": 1, "february": 2, "feb": 2, "march": 3, "mar": 3,
    "april": 4, "apr": 4, "may": 5, "june": 6, "jun": 6, "july": 7, "jul": 7,
    "august": 8, "aug": 8, "september": 9, "sep": 9, "sept": 9, "october": 10,
    "oct": 10, "november": 11, "nov": 11, "december": 12, "dec": 12,
}


def detect_date(text: str) -> Optional[str]:
    """Return the most plausible document date as an ISO ``YYYY-MM-DD`` string.

    Heuristic: collect every parseable date, then prefer the *latest* date that
    is not in the future (statements/invoices are usually dated at issue). This
    avoids picking, e.g., an account-opened date from years ago.
    """
    if not text:
        return None

    candidates: List[date] = []
    today = date.today()

    for pattern, kind in DATE_PATTERNS:
        for m in pattern.finditer(text):
            try:
                if kind == "ymd":
                    y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
                elif kind == "mdy":
                    mo, d, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
                elif kind == "month_name":
                    mo = _MONTH_INDEX[m.group(1).lower()]
                    d, y = int(m.group(2)), int(m.group(3))
                elif kind == "day_month_name":
                    d = int(m.group(1))
                    mo = _MONTH_INDEX[m.group(2).lower()]
                    y = int(m.group(3))
                else:
                    continue
                cand = date(y, mo, d)
            except (ValueError, KeyError):
                continue
            # Reject obviously wrong dates.
            if cand.year < 1990 or cand > today:
                continue
            candidates.append(cand)

    if not candidates:
        return None
    best = max(candidates)
    return best.isoformat()


def detect_bank(text: str) -> Optional[str]:
    """Return the first known bank/issuer mentioned, normalized for folders."""
    if not text:
        return None
    lowered = text.lower()
    # Normalize a few aliases to a canonical display name.
    aliases = {
        "jpmorgan": "Chase", "amex": "American Express",
        "u.s. bank": "US Bank", "citi": "Citibank",
    }
    for bank in KNOWN_BANKS:
        if bank.lower() in lowered:
            return aliases.get(bank.lower(), bank)
    return None


# Company/party detection: look for common "labeled" lines first, else fall
# back to a capitalized-name heuristic near the top of the document.
_PARTY_LABELS = [
    re.compile(r"(?:bill to|billed to|sold to|customer)\s*[:\-]\s*(.+)", re.I),
    re.compile(r"(?:from|vendor|company|payee|remit to)\s*[:\-]\s*(.+)", re.I),
    re.compile(r"(?:name)\s*[:\-]\s*(.+)", re.I),
]


def detect_party(text: str, bank: Optional[str]) -> Optional[str]:
    """Best-effort detection of the counterparty / company / person."""
    if bank:
        return bank
    if not text:
        return None
    head = text[:2000]
    for pat in _PARTY_LABELS:
        m = pat.search(head)
        if m:
            value = m.group(1).strip()
            value = re.split(r"[\n\r]", value)[0].strip()
            value = re.sub(r"\s{2,}", " ", value)
            if 2 <= len(value) <= 60:
                return _clean_token(value)
    # Fallback: first line that looks like a proper name / company.
    for line in head.splitlines():
        line = line.strip()
        if 3 <= len(line) <= 50 and re.match(
            r"^[A-Z][A-Za-z0-9&.,'\- ]+$", line
        ):
            # Avoid all-caps banner noise like "STATEMENT OF ACCOUNT".
            words = line.split()
            if 1 <= len(words) <= 6 and not line.isupper():
                return _clean_token(line)
    return None


# --------------------------------------------------------------------------- #
#  Filename helpers
# --------------------------------------------------------------------------- #

_INVALID_FS_CHARS = re.compile(r'[<>:"/\\|?*\x00-\x1f]')


def _clean_token(value: str) -> str:
    """Make a string safe and tidy for use inside a filename."""
    value = _INVALID_FS_CHARS.sub("", value)
    value = value.replace("\t", " ").strip(" .-")
    value = re.sub(r"\s{2,}", " ", value)
    return value[:60].strip()


def sanitize_filename(name: str) -> str:
    """Return a filesystem-safe version of ``name`` (no path separators)."""
    name = _INVALID_FS_CHARS.sub("_", name)
    name = name.strip().strip(".")
    return name or "unnamed"


def unique_destination(dest_dir: Path, filename: str) -> Path:
    """Return a non-colliding path in ``dest_dir`` for ``filename``.

    Never overwrites: if ``report.pdf`` exists we return ``report (2).pdf``,
    then ``report (3).pdf``, and so on.
    """
    candidate = dest_dir / filename
    if not candidate.exists():
        return candidate
    stem = candidate.stem
    suffix = candidate.suffix
    n = 2
    while True:
        candidate = dest_dir / f"{stem} ({n}){suffix}"
        if not candidate.exists():
            return candidate
        n += 1


def build_new_name(
    original_name: str,
    category: str,
    detected_date: Optional[str],
    party: Optional[str],
    ext: str,
) -> str:
    """Build the standardized name:

        ``YYYY-MM-DD - Category - Party - Original Hint.ext``

    Missing parts are skipped gracefully so we never produce ugly
    ``  -  -  `` runs.
    """
    hint = Path(original_name).stem
    hint = _clean_token(hint)[:40]

    parts: List[str] = []
    if detected_date:
        parts.append(detected_date)
    parts.append(category)
    if party:
        parts.append(party)
    if hint and hint.lower() not in {p.lower() for p in parts}:
        parts.append(hint)

    base = " - ".join(_clean_token(p) for p in parts if p)
    return sanitize_filename(base) + ext.lower()


# --------------------------------------------------------------------------- #
#  Text extraction (with graceful degradation + OCR)
# --------------------------------------------------------------------------- #

class TextExtractor:
    """Extracts text from supported file types.

    Optional dependencies are imported lazily. If a dependency is missing we
    log a one-time warning and return an empty string for that type, so the
    program keeps working for everything else.
    """

    def __init__(self, use_ocr: bool = False, scanned_pdf_ocr: bool = False,
                 ocr_lang: str = "eng"):
        self.use_ocr = use_ocr
        self.scanned_pdf_ocr = scanned_pdf_ocr
        self.ocr_lang = ocr_lang
        self._warned: set = set()

    # -- small helper to warn only once per missing dependency -------------- #
    def _warn_once(self, key: str, message: str) -> None:
        if key not in self._warned:
            self._warned.add(key)
            logger.warning(message)

    # -- public entry point ------------------------------------------------- #
    def extract(self, path: Path) -> Tuple[str, bool]:
        """Return ``(text, used_ocr)`` for the file at ``path``.

        Never raises: extraction failures are logged and yield ``("", False)``.
        """
        ext = path.suffix.lower()
        try:
            if ext == ".pdf":
                return self._extract_pdf(path)
            if ext == ".docx":
                return self._extract_docx(path), False
            if ext == ".doc":
                return self._extract_doc(path), False
            if ext in (".xlsx", ".xls"):
                return self._extract_excel(path), False
            if ext == ".csv":
                return self._extract_text_file(path), False
            if ext == ".txt":
                return self._extract_text_file(path), False
            if ext in IMAGE_EXTENSIONS:
                return self._extract_image(path)
            if ext == ".eml":
                return self._extract_eml(path), False
            if ext == ".msg":
                return self._extract_msg(path), False
            if ext == ".zip":
                return self._extract_zip(path)
        except Exception as exc:  # noqa: BLE001 - we truly want to keep going
            logger.debug("Extraction failed for %s: %s", path, exc)
        return "", False

    # -- PDF ---------------------------------------------------------------- #
    def _extract_pdf(self, path: Path) -> Tuple[str, bool]:
        text = ""
        # First try a normal text extraction (pdfplumber, then PyPDF2).
        try:
            import pdfplumber  # type: ignore

            with pdfplumber.open(str(path)) as pdf:
                pages = []
                for page in pdf.pages[:30]:  # cap pages for speed
                    pages.append(page.extract_text() or "")
                text = "\n".join(pages)
        except ImportError:
            try:
                from PyPDF2 import PdfReader  # type: ignore

                reader = PdfReader(str(path))
                pages = [(p.extract_text() or "") for p in reader.pages[:30]]
                text = "\n".join(pages)
            except ImportError:
                self._warn_once(
                    "pdf",
                    "No PDF library found (install pdfplumber or PyPDF2); "
                    "PDF text extraction disabled.",
                )
            except Exception as exc:  # noqa: BLE001
                logger.debug("PyPDF2 failed on %s: %s", path, exc)
        except Exception as exc:  # noqa: BLE001
            logger.debug("pdfplumber failed on %s: %s", path, exc)

        # If we got little/no text and the user allowed scanned-PDF OCR, OCR it.
        if self.scanned_pdf_ocr and len(text.strip()) < 40:
            ocr_text = self._ocr_pdf(path)
            if ocr_text.strip():
                return ocr_text, True
        return text, False

    def _ocr_pdf(self, path: Path) -> str:
        """Rasterize a PDF and OCR each page. Requires pdf2image + pytesseract
        (and Poppler installed on the system for pdf2image)."""
        try:
            from pdf2image import convert_from_path  # type: ignore
            import pytesseract  # type: ignore
        except ImportError:
            self._warn_once(
                "pdf_ocr",
                "Scanned-PDF OCR needs pdf2image + pytesseract (and Poppler "
                "+ Tesseract installed); skipping OCR of scanned PDFs.",
            )
            return ""
        try:
            images = convert_from_path(str(path), dpi=200)
        except Exception as exc:  # noqa: BLE001
            self._warn_once(
                "poppler",
                "pdf2image could not rasterize PDFs (is Poppler installed and "
                "on PATH?). Scanned-PDF OCR disabled. Detail: %s" % exc,
            )
            return ""
        chunks = []
        for img in images[:15]:  # cap for speed
            try:
                chunks.append(pytesseract.image_to_string(img,
                                                           lang=self.ocr_lang))
            except Exception as exc:  # noqa: BLE001
                logger.debug("Tesseract failed on a PDF page of %s: %s",
                             path, exc)
        return "\n".join(chunks)

    # -- DOCX / DOC --------------------------------------------------------- #
    def _extract_docx(self, path: Path) -> str:
        try:
            import docx  # type: ignore
        except ImportError:
            self._warn_once(
                "docx",
                "python-docx not installed; DOCX extraction disabled.",
            )
            return ""
        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" ".join(c.text for c in row.cells))
        return "\n".join(parts)

    def _extract_doc(self, path: Path) -> str:
        # Legacy .doc is binary; best-effort. Try textract if present.
        try:
            import textract  # type: ignore

            return textract.process(str(path)).decode("utf-8", "ignore")
        except Exception:  # noqa: BLE001
            self._warn_once(
                "doc",
                "Legacy .doc extraction needs 'textract' (and antiword); "
                ".doc files will be sent to Needs Review.",
            )
            return ""

    # -- Excel -------------------------------------------------------------- #
    def _extract_excel(self, path: Path) -> str:
        ext = path.suffix.lower()
        try:
            if ext == ".xlsx":
                import openpyxl  # type: ignore

                wb = openpyxl.load_workbook(str(path), read_only=True,
                                            data_only=True)
                parts = []
                for ws in wb.worksheets:
                    for row in ws.iter_rows(values_only=True):
                        parts.append(
                            " ".join("" if c is None else str(c) for c in row)
                        )
                        if len(parts) > 2000:
                            break
                return "\n".join(parts)
            else:  # .xls
                import xlrd  # type: ignore

                book = xlrd.open_workbook(str(path))
                parts = []
                for sheet in book.sheets():
                    for r in range(min(sheet.nrows, 2000)):
                        parts.append(
                            " ".join(str(c.value) for c in sheet.row(r))
                        )
                return "\n".join(parts)
        except ImportError:
            self._warn_once(
                "excel",
                "Excel extraction needs openpyxl (.xlsx) / xlrd (.xls); "
                "spreadsheet extraction disabled.",
            )
            return ""

    # -- Plain text / CSV --------------------------------------------------- #
    def _extract_text_file(self, path: Path) -> str:
        for enc in ("utf-8", "utf-16", "latin-1"):
            try:
                with open(path, "r", encoding=enc, errors="strict") as fh:
                    return fh.read(MAX_TEXT_FOR_CLASSIFY * 2)
            except (UnicodeError, UnicodeDecodeError):
                continue
            except Exception as exc:  # noqa: BLE001
                logger.debug("Text read failed for %s: %s", path, exc)
                return ""
        # Last resort: read bytes and decode loosely.
        try:
            return path.read_bytes()[: MAX_TEXT_FOR_CLASSIFY * 2].decode(
                "utf-8", "ignore"
            )
        except Exception:  # noqa: BLE001
            return ""

    # -- Images (OCR) ------------------------------------------------------- #
    def _extract_image(self, path: Path) -> Tuple[str, bool]:
        if not self.use_ocr:
            return "", False
        try:
            from PIL import Image  # type: ignore
            import pytesseract  # type: ignore
        except ImportError:
            self._warn_once(
                "image_ocr",
                "Image OCR needs Pillow + pytesseract (and Tesseract "
                "installed); image text extraction disabled.",
            )
            return "", False
        try:
            with Image.open(path) as img:
                text = pytesseract.image_to_string(img, lang=self.ocr_lang)
            return text, True
        except Exception as exc:  # noqa: BLE001
            self._warn_once(
                "tesseract",
                "Tesseract OCR failed (is the Tesseract engine installed and "
                "on PATH?). Detail: %s" % exc,
            )
            return "", False

    # -- Email .eml --------------------------------------------------------- #
    def _extract_eml(self, path: Path) -> str:
        import email
        from email import policy

        with open(path, "rb") as fh:
            msg = email.message_from_binary_file(fh, policy=policy.default)
        parts = [
            f"Subject: {msg.get('subject', '')}",
            f"From: {msg.get('from', '')}",
            f"To: {msg.get('to', '')}",
            f"Date: {msg.get('date', '')}",
        ]
        try:
            body = msg.get_body(preferencelist=("plain", "html"))
            if body is not None:
                content = body.get_content()
                # crude HTML strip
                content = re.sub(r"<[^>]+>", " ", content)
                parts.append(content)
        except Exception as exc:  # noqa: BLE001
            logger.debug("EML body parse failed for %s: %s", path, exc)
        return "\n".join(parts)

    # -- Outlook .msg ------------------------------------------------------- #
    def _extract_msg(self, path: Path) -> str:
        try:
            import extract_msg  # type: ignore
        except ImportError:
            self._warn_once(
                "msg",
                "Outlook .msg extraction needs 'extract-msg'; .msg files go "
                "to Needs Review.",
            )
            return ""
        try:
            m = extract_msg.Message(str(path))
            return "\n".join(
                filter(None, [
                    f"Subject: {m.subject or ''}",
                    f"From: {m.sender or ''}",
                    f"To: {m.to or ''}",
                    f"Date: {m.date or ''}",
                    m.body or "",
                ])
            )
        except Exception as exc:  # noqa: BLE001
            logger.debug("extract_msg failed for %s: %s", path, exc)
            return ""

    # -- ZIP (classify by contained text) ----------------------------------- #
    def _extract_zip(self, path: Path) -> Tuple[str, bool]:
        """Peek inside a ZIP: gather names + a little text from inner text-ish
        files so the archive can be classified by what it contains."""
        parts: List[str] = []
        used_ocr = False
        try:
            with zipfile.ZipFile(path) as zf:
                names = zf.namelist()
                parts.append("ARCHIVE CONTENTS: " + " ".join(names[:100]))
                for name in names:
                    inner_ext = os.path.splitext(name)[1].lower()
                    if inner_ext in (".txt", ".csv"):
                        try:
                            with zf.open(name) as fh:
                                data = fh.read(50000)
                            parts.append(data.decode("utf-8", "ignore"))
                        except Exception:  # noqa: BLE001
                            continue
                    if sum(len(p) for p in parts) > MAX_TEXT_FOR_CLASSIFY:
                        break
        except zipfile.BadZipFile:
            logger.debug("Bad ZIP: %s", path)
        return "\n".join(parts), used_ocr


# --------------------------------------------------------------------------- #
#  Classifier
# --------------------------------------------------------------------------- #

@dataclass
class Classification:
    category: str
    confidence: float
    needs_review: bool
    reason: str  # human-readable score breakdown
    scores: Dict[str, float] = field(default_factory=dict)


def _phrase_present(text_lower: str, phrase) -> bool:
    """Return True if ``phrase`` (str, compiled regex, or tuple of either)
    appears in ``text_lower``. A tuple is an OR-group (any match counts)."""
    if isinstance(phrase, tuple):
        return any(_phrase_present(text_lower, p) for p in phrase)
    if isinstance(phrase, re.Pattern):
        return bool(phrase.search(text_lower))
    return phrase.lower() in text_lower


class Classifier:
    """Rule-based, transparent, confidence-scoring classifier."""

    def __init__(self, rules: Dict[str, Rule], min_confidence: float = 0.45,
                 corrections: Optional[Dict[str, str]] = None):
        self.rules = rules
        self.min_confidence = min_confidence
        # sha256 -> corrected category (from --learn-from)
        self.corrections = corrections or {}

    def classify(self, text: str, sha256: str,
                 filename: str = "") -> Classification:
        # 1) Learned corrections always win (exact content match by hash).
        if sha256 in self.corrections:
            cat = self.corrections[sha256]
            return Classification(
                category=cat,
                confidence=1.0,
                needs_review=False,
                reason="Matched a learned correction by SHA-256.",
                scores={cat: CONFIDENCE_FULL_SCALE},
            )

        # We classify on extracted text; the filename is added as a light hint.
        blob = (text or "")[:MAX_TEXT_FOR_CLASSIFY]
        if filename:
            blob = blob + "\n" + filename
        text_lower = blob.lower()

        if len(text_lower.strip()) < 15:
            # No usable text at all -> cannot responsibly classify.
            return Classification(
                category=CAT_NEEDS_REVIEW,
                confidence=0.0,
                needs_review=True,
                reason="No extractable text (possibly scanned without OCR, "
                       "empty, or an unsupported binary).",
            )

        scores: Dict[str, float] = {}
        breakdowns: Dict[str, str] = {}

        for name, rule in self.rules.items():
            score, breakdown = self._score_rule(text_lower, rule)
            scores[name] = score
            breakdowns[name] = breakdown

        # Pick the best-scoring category.
        best_cat = max(scores, key=lambda k: scores[k])
        best_raw = scores[best_cat]
        confidence = max(0.0, min(1.0, best_raw / CONFIDENCE_FULL_SCALE))

        # Margin check: if the runner-up is basically tied, we're not confident.
        ordered = sorted(scores.values(), reverse=True)
        margin = ordered[0] - (ordered[1] if len(ordered) > 1 else 0.0)

        needs_review = (
            best_raw <= 0.0
            or confidence < self.min_confidence
            or margin < 1.0  # ambiguous between two categories
        )

        final_cat = best_cat if not needs_review else CAT_NEEDS_REVIEW

        reason = (
            f"best={best_cat} raw={best_raw:.2f} conf={confidence:.2f} "
            f"margin={margin:.2f} thr={self.min_confidence:.2f} | "
            f"{breakdowns.get(best_cat, '')}"
        )

        return Classification(
            category=final_cat,
            confidence=round(confidence, 3),
            needs_review=needs_review,
            reason=reason,
            scores={k: round(v, 2) for k, v in scores.items() if v > 0},
        )

    def _score_rule(self, text_lower: str, rule: Rule) -> Tuple[float, str]:
        """Score one category. Returns ``(score, breakdown_string)``."""
        # Must-have gate.
        must_hits = sum(
            1 for group in rule.must_have if _phrase_present(text_lower, group)
        )
        if rule.must_have and must_hits < rule.min_must_have:
            return 0.0, (
                f"disqualified: only {must_hits}/{rule.min_must_have} "
                f"must-have signals"
            )

        score = rule.base + must_hits * MUST_HAVE_WEIGHT

        strong_hits = [s for s in rule.strong if _phrase_present(text_lower, s)]
        weak_hits = [w for w in rule.weak if _phrase_present(text_lower, w)]
        neg_hits = [n for n in rule.negative
                    if _phrase_present(text_lower, n)]

        score += len(strong_hits) * STRONG_WEIGHT
        score += len(weak_hits) * WEAK_WEIGHT
        score -= len(neg_hits) * NEG_WEIGHT
        score = max(0.0, score)

        breakdown = (
            f"must={must_hits} strong={len(strong_hits)} "
            f"weak={len(weak_hits)} neg={len(neg_hits)}"
        )
        return score, breakdown


# --------------------------------------------------------------------------- #
#  Duplicate detection
# --------------------------------------------------------------------------- #

def sha256_of_file(path: Path, chunk_size: int = 1 << 20) -> str:
    """Compute the SHA-256 of a file, streaming so big files are fine."""
    h = hashlib.sha256()
    with open(path, "rb") as fh:
        for chunk in iter(lambda: fh.read(chunk_size), b""):
            h.update(chunk)
    return h.hexdigest()


# --------------------------------------------------------------------------- #
#  Record + persistence (CSV + SQLite)
# --------------------------------------------------------------------------- #

# The canonical, ordered set of columns for both the CSV and the DB.
FIELDNAMES = [
    "original_path",
    "original_name",
    "original_size",
    "new_path",
    "new_name",
    "category",
    "confidence",
    "needs_review",
    "detected_date",
    "detected_party",
    "detected_bank",
    "sha256",
    "duplicate_group",
    "duplicate_status",
    "duplicate_of",
    "reason",           # a.k.a. score_details
    "text_preview",
    "corrected_category",
    "status",
]


@dataclass
class Record:
    original_path: str = ""
    original_name: str = ""
    original_size: int = 0
    new_path: str = ""
    new_name: str = ""
    category: str = ""
    confidence: float = 0.0
    needs_review: bool = False
    detected_date: str = ""
    detected_party: str = ""
    detected_bank: str = ""
    sha256: str = ""
    duplicate_group: str = ""
    duplicate_status: str = ""   # original | duplicate | ""
    duplicate_of: str = ""
    reason: str = ""
    text_preview: str = ""
    corrected_category: str = ""  # filled by the user, read by --learn-from
    status: str = ""             # planned | copied | moved | skipped | error

    def as_row(self) -> Dict[str, str]:
        d = asdict(self)
        # Booleans -> ints for tidy CSV/DB storage.
        d["needs_review"] = int(bool(self.needs_review))
        return d


class Store:
    """Writes records to both a CSV log and an SQLite database."""

    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
        self.csv_path = output_dir / LOG_CSV_NAME
        self.db_path = output_dir / DB_NAME
        self._csv_file = None
        self._csv_writer = None
        self._conn: Optional[sqlite3.Connection] = None

    def open(self) -> None:
        self.output_dir.mkdir(parents=True, exist_ok=True)
        # CSV (UTF-8 with BOM so Excel on Windows opens it cleanly).
        self._csv_file = open(
            self.csv_path, "w", newline="", encoding="utf-8-sig"
        )
        self._csv_writer = csv.DictWriter(self._csv_file, fieldnames=FIELDNAMES)
        self._csv_writer.writeheader()
        # SQLite.
        self._conn = sqlite3.connect(str(self.db_path))
        cols = ", ".join(f"{c} TEXT" for c in FIELDNAMES)
        self._conn.execute(f"CREATE TABLE IF NOT EXISTS files ({cols})")
        # Fresh run: clear previous rows so the DB mirrors this run's CSV.
        self._conn.execute("DELETE FROM files")
        # PERSISTENT history: one row per unique content hash ever organized
        # by an --apply run into this output folder. This table is NEVER
        # wiped, which is what gives us duplicate detection ACROSS runs
        # (re-running on the same source won't re-copy files, and new files
        # that duplicate already-organized content are flagged).
        self._conn.execute(
            "CREATE TABLE IF NOT EXISTS history ("
            "  sha256 TEXT PRIMARY KEY,"
            "  original_path TEXT,"
            "  new_path TEXT,"
            "  category TEXT,"
            "  organized_at TEXT"
            ")"
        )
        self._conn.commit()

    def load_history(self) -> Dict[str, Dict[str, str]]:
        """Return ``{sha256: {original_path, new_path, category}}`` for every
        file organized by a previous --apply run into this output folder."""
        if self._conn is None:
            return {}
        rows = self._conn.execute(
            "SELECT sha256, original_path, new_path, category FROM history"
        ).fetchall()
        return {
            r[0]: {"original_path": r[1], "new_path": r[2], "category": r[3]}
            for r in rows
        }

    def record_history(self, rec: "Record") -> None:
        """Remember that this content hash has been organized (apply runs
        only), so future runs can skip it instead of duplicating output."""
        if self._conn is None:
            return
        self._conn.execute(
            "INSERT OR REPLACE INTO history "
            "(sha256, original_path, new_path, category, organized_at) "
            "VALUES (?, ?, ?, ?, ?)",
            (rec.sha256, rec.original_path, rec.new_path, rec.category,
             datetime.now().isoformat(timespec="seconds")),
        )

    def write(self, record: Record) -> None:
        row = record.as_row()
        if self._csv_writer is not None:
            self._csv_writer.writerow(row)
            self._csv_file.flush()
        if self._conn is not None:
            placeholders = ", ".join("?" for _ in FIELDNAMES)
            self._conn.execute(
                f"INSERT INTO files ({', '.join(FIELDNAMES)}) "
                f"VALUES ({placeholders})",
                [str(row[c]) for c in FIELDNAMES],
            )

    def close(self) -> None:
        if self._conn is not None:
            self._conn.commit()
            self._conn.close()
            self._conn = None
        if self._csv_file is not None:
            self._csv_file.close()
            self._csv_file = None


# --------------------------------------------------------------------------- #
#  Learning from corrections
# --------------------------------------------------------------------------- #

def load_corrections(csv_path: Path) -> Dict[str, str]:
    """Read a previously-exported CSV that the user edited, and return a
    mapping of ``sha256 -> corrected_category``.

    A correction is any row where ``corrected_category`` is non-empty and
    differs from ``category``. Keyed by SHA-256, so it applies to the exact
    same file content wherever it appears next time.
    """
    corrections: Dict[str, str] = {}
    if not csv_path.exists():
        logger.error("Correction file not found: %s", csv_path)
        return corrections
    with open(csv_path, "r", encoding="utf-8-sig", newline="") as fh:
        reader = csv.DictReader(fh)
        for row in reader:
            sha = (row.get("sha256") or "").strip()
            corrected = (row.get("corrected_category") or "").strip()
            if sha and corrected:
                corrections[sha] = corrected
    logger.info("Loaded %d correction(s) from %s", len(corrections), csv_path)
    return corrections


# --------------------------------------------------------------------------- #
#  The organizer engine
# --------------------------------------------------------------------------- #

@dataclass
class Options:
    folder: Path
    output: Path
    apply: bool = False
    copy: bool = True            # copy is the safe default operation
    use_ocr: bool = False
    scanned_pdf_ocr: bool = False
    duplicate_action: str = "folder"   # keep | folder | skip
    min_confidence: float = 0.45
    move_needs_review: bool = False
    # Inside _Needs_Review, group files into subfolders by their file type
    # (extension), e.g. _Needs_Review/PDF, _Needs_Review/MP3, _Needs_Review/ZIP.
    review_by_type: bool = False
    rename: bool = False
    learn_from: Optional[Path] = None
    ocr_lang: str = "eng"
    # Ignore the persistent history and process everything again.
    reprocess: bool = False


class Organizer:
    """Orchestrates scanning, extraction, classification, dedup, and the
    (optional) copy/move actions."""

    def __init__(self, options: Options,
                 progress: Optional[Callable[[int, int, Record], None]] = None):
        self.opt = options
        self.progress = progress
        corrections = (
            load_corrections(options.learn_from) if options.learn_from else {}
        )
        self.extractor = TextExtractor(
            use_ocr=options.use_ocr,
            scanned_pdf_ocr=options.scanned_pdf_ocr,
            ocr_lang=options.ocr_lang,
        )
        self.classifier = Classifier(
            RULES, min_confidence=options.min_confidence,
            corrections=corrections,
        )
        self.store = Store(options.output)
        # sha256 -> Record of the FIRST (original) file we saw with that hash
        # in THIS run.
        self._seen_hashes: Dict[str, Record] = {}
        # sha256 -> info for content organized by PREVIOUS --apply runs into
        # this same output folder (loaded from the DB's history table).
        self._prior_hashes: Dict[str, Dict[str, str]] = {}
        self.records: List[Record] = []
        self._output_resolved = options.output.resolve()

    # -- scanning ----------------------------------------------------------- #
    def _iter_files(self) -> List[Path]:
        """Walk the source tree, skipping system folders and, crucially, the
        output folder (so we never re-organize our own output)."""
        found: List[Path] = []
        for root, dirs, files in os.walk(self.opt.folder):
            root_path = Path(root)
            # Prune: skip system dirs and the output tree.
            pruned = []
            for d in dirs:
                lower = d.lower()
                child = (root_path / d).resolve()
                if lower in SKIP_DIR_NAMES:
                    continue
                if child == self._output_resolved or self._is_within_output(
                    child
                ):
                    continue
                pruned.append(d)
            dirs[:] = pruned

            for fname in files:
                p = root_path / fname
                if p.name in (LOG_CSV_NAME, DB_NAME):
                    continue
                found.append(p)
        return found

    def _is_within_output(self, path: Path) -> bool:
        try:
            path.resolve().relative_to(self._output_resolved)
            return True
        except (ValueError, OSError):
            return False

    # -- main run ----------------------------------------------------------- #
    def run(self) -> List[Record]:
        self.store.open()
        try:
            # Load cross-run memory unless the user asked to reprocess all.
            if not self.opt.reprocess:
                self._prior_hashes = self.store.load_history()
                if self._prior_hashes:
                    logger.info(
                        "Loaded %d previously organized file hash(es) from "
                        "the database (use --reprocess to ignore).",
                        len(self._prior_hashes),
                    )
            files = self._iter_files()
            total = len(files)
            logger.info("Scanning %d file(s) under %s", total, self.opt.folder)
            for i, path in enumerate(files, start=1):
                record = self._process_file(path)
                self.records.append(record)
                self.store.write(record)
                # Remember successfully organized content for future runs.
                if self.opt.apply and record.status in ("copied", "moved"):
                    self.store.record_history(record)
                if self.progress:
                    try:
                        self.progress(i, total, record)
                    except Exception:  # noqa: BLE001 - GUI callback safety
                        pass
        finally:
            self.store.close()
        self._log_summary()
        return self.records

    # -- per-file pipeline -------------------------------------------------- #
    def _process_file(self, path: Path) -> Record:
        rec = Record()
        rec.original_path = str(path)
        rec.original_name = path.name
        rec.status = "planned"

        # File stats + hash first (needed for dedup + learning).
        try:
            rec.original_size = path.stat().st_size
            rec.sha256 = sha256_of_file(path)
        except Exception as exc:  # noqa: BLE001
            rec.status = "error"
            rec.reason = f"Could not read file: {exc}"
            rec.category = CAT_NEEDS_REVIEW
            rec.needs_review = True
            return rec

        rec.duplicate_group = rec.sha256[:12]

        # -- cross-run duplicate check (persistent history) ------------------ #
        # If this exact content was already organized by a previous --apply
        # run into this output folder, don't copy/move it again: log it as
        # already organized. This makes re-runs incremental and prevents
        # " (2)" copies from piling up in the output tree.
        if rec.sha256 in self._prior_hashes:
            prior = self._prior_hashes[rec.sha256]
            rec.category = prior.get("category") or CAT_DUPLICATE
            rec.duplicate_status = "already organized"
            rec.duplicate_of = prior.get("new_path") or \
                prior.get("original_path") or ""
            rec.reason = ("Identical content was already organized in a "
                          "previous run (SHA-256 match in history).")
            rec.status = "skipped (already organized in a previous run)"
            return rec

        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            rec.category = CAT_NEEDS_REVIEW
            rec.needs_review = True
            rec.reason = f"Unsupported file type '{ext}'."
            rec.duplicate_status = self._dedupe_status(rec)
            self._plan_and_act(path, rec)
            return rec

        # -- duplicate detection (by content hash) -------------------------- #
        is_duplicate = rec.sha256 in self._seen_hashes
        rec.duplicate_status = self._dedupe_status(rec)
        if is_duplicate:
            original = self._seen_hashes[rec.sha256]
            # Link to where the original ENDED UP in the organized tree
            # (its new_path), falling back to its source path if it wasn't
            # moved/copied (e.g. needs-review files that were only logged).
            # Duplicates are always processed after their original, so the
            # original's destination is already known at this point.
            rec.duplicate_of = original.new_path or original.original_path
            # Carry the original's classification so logs are coherent.
            rec.category = original.category
            rec.confidence = original.confidence
            rec.detected_date = original.detected_date
            rec.detected_party = original.detected_party
            rec.detected_bank = original.detected_bank
            rec.needs_review = original.needs_review
            rec.reason = (f"Exact duplicate of {original.original_name} "
                          f"(original kept at: "
                          f"{original.new_path or original.original_path}).")
            rec.text_preview = original.text_preview
            self._handle_duplicate(path, rec, original)
            return rec

        # -- first time we've seen this content: classify it ---------------- #
        text, used_ocr = self.extractor.extract(path)
        rec.text_preview = _make_preview(text)

        result = self.classifier.classify(text, rec.sha256, filename=path.name)
        rec.category = result.category
        rec.confidence = result.confidence
        rec.needs_review = result.needs_review
        rec.reason = result.reason + (" [OCR]" if used_ocr else "")

        # Field detection (date / bank / party) for renaming + foldering.
        rec.detected_date = detect_date(text) or ""
        rec.detected_bank = detect_bank(text) or ""
        rec.detected_party = detect_party(text, detect_bank(text)) or ""

        # Remember this as the original for its hash BEFORE acting.
        self._seen_hashes[rec.sha256] = rec

        self._plan_and_act(path, rec)
        return rec

    def _dedupe_status(self, rec: Record) -> str:
        return "duplicate" if rec.sha256 in self._seen_hashes else "original"

    # -- duplicate handling per the chosen policy ---------------------------- #
    def _handle_duplicate(self, path: Path, rec: Record,
                          original: Record) -> None:
        action = self.opt.duplicate_action
        if action == "skip":
            # Leave the file exactly where it is; only log it.
            rec.new_path = ""
            rec.new_name = ""
            rec.status = "skipped (duplicate: left in place)"
            return
        if action == "keep":
            # Organize it normally like any other file (but never overwrite).
            self._plan_and_act(path, rec)
            if rec.status.startswith("planned"):
                rec.status = "planned (duplicate: organized normally)"
            return
        # Default "folder": route duplicates into the _Duplicates tree.
        dest_dir = self.opt.output / CATEGORY_FOLDER[CAT_DUPLICATE]
        self._finalize_destination(path, rec, dest_dir)

    # -- decide destination folder + name, then (optionally) act ------------ #
    def _plan_and_act(self, path: Path, rec: Record) -> None:
        # Files that need review either go to _Needs_Review or are only logged.
        if rec.needs_review and not self.opt.move_needs_review:
            rec.new_path = ""
            rec.new_name = ""
            rec.status = "logged only (needs review, not moved)"
            return

        dest_dir = self._destination_dir(rec)
        self._finalize_destination(path, rec, dest_dir)

    def _destination_dir(self, rec: Record) -> Path:
        """Compute the target directory (including Bank/Year sub-levels)."""
        category = rec.category if rec.category in CATEGORY_FOLDER \
            else CAT_NEEDS_REVIEW
        rel = CATEGORY_FOLDER[category]
        dest = self.opt.output / rel

        # Add /<Bank>/<Year> depth for bank & credit-card statements when known.
        if category in ("Bank Statements", "Credit Card Statements"):
            if rec.detected_bank:
                dest = dest / sanitize_filename(rec.detected_bank)
            if rec.detected_date[:4].isdigit():
                dest = dest / rec.detected_date[:4]
        # Add /<Party> for invoices/receipts/utilities when known.
        elif category in ("Invoices", "Receipts", "Utilities") \
                and rec.detected_party:
            dest = dest / sanitize_filename(rec.detected_party)
        # Optionally group needs-review files by file type so they're easy
        # to triage later: _Needs_Review/PDF, _Needs_Review/MP3, ...
        elif category == CAT_NEEDS_REVIEW and self.opt.review_by_type:
            ext = Path(rec.original_name).suffix.lstrip(".").upper()
            dest = dest / sanitize_filename(ext or "No Extension")
        return dest

    def _finalize_destination(self, path: Path, rec: Record,
                              dest_dir: Path) -> None:
        # Choose the filename (renamed or original).
        if self.opt.rename and not rec.needs_review:
            new_name = build_new_name(
                rec.original_name, rec.category, rec.detected_date or None,
                rec.detected_party or None, path.suffix,
            )
        else:
            new_name = rec.original_name

        if self.opt.apply:
            dest_dir.mkdir(parents=True, exist_ok=True)
            dest = unique_destination(dest_dir, new_name)
            rec.new_path = str(dest)
            rec.new_name = dest.name
            self._perform_io(path, dest, rec)
        else:
            # Dry-run: compute where it *would* go (still collision-aware,
            # but without touching the filesystem).
            dest = self._dry_run_dest(dest_dir, new_name)
            rec.new_path = str(dest)
            rec.new_name = dest.name
            op = "copy" if self.opt.copy else "move"
            rec.status = f"planned ({op}, dry-run)"

    def _dry_run_dest(self, dest_dir: Path, filename: str) -> Path:
        """Collision-aware destination for dry-run, tracking names we've
        already planned this run so two files don't 'both' become the same
        target on paper."""
        candidate = dest_dir / filename
        planned = getattr(self, "_planned_names", None)
        if planned is None:
            planned = set()
            self._planned_names = planned
        stem, suffix = Path(filename).stem, Path(filename).suffix
        n = 2
        while str(candidate).lower() in planned or candidate.exists():
            candidate = dest_dir / f"{stem} ({n}){suffix}"
            n += 1
        planned.add(str(candidate).lower())
        return candidate

    def _perform_io(self, src: Path, dest: Path, rec: Record) -> None:
        """Actually copy or move, safely. Never overwrites, never deletes the
        source in copy mode."""
        try:
            if self.opt.copy:
                shutil.copy2(str(src), str(dest))
                rec.status = "copied"
            else:
                shutil.move(str(src), str(dest))
                rec.status = "moved"
        except Exception as exc:  # noqa: BLE001
            rec.status = f"error: {exc}"
            logger.error("I/O failed for %s -> %s: %s", src, dest, exc)

    # -- summary ------------------------------------------------------------ #
    def _log_summary(self) -> None:
        by_cat: Dict[str, int] = {}
        dupes = 0
        review = 0
        for r in self.records:
            by_cat[r.category] = by_cat.get(r.category, 0) + 1
            if r.duplicate_status == "duplicate":
                dupes += 1
            if r.needs_review:
                review += 1
        logger.info("-" * 60)
        logger.info("Processed %d file(s).", len(self.records))
        logger.info("Duplicates: %d | Needs review: %d", dupes, review)
        for cat in sorted(by_cat):
            logger.info("  %-26s %d", cat, by_cat[cat])
        logger.info("Log CSV : %s", self.store.csv_path)
        logger.info("Database: %s", self.store.db_path)
        if not self.opt.apply:
            logger.info("DRY-RUN complete. Re-run with --apply to make changes.")
        logger.info("-" * 60)


def _make_preview(text: str) -> str:
    if not text:
        return ""
    collapsed = re.sub(r"\s+", " ", text).strip()
    return collapsed[:TEXT_PREVIEW_LEN]


# --------------------------------------------------------------------------- #
#  Command-line interface
# --------------------------------------------------------------------------- #

def build_arg_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="smart_document_organizer",
        description=f"{APP_NAME} v{APP_VERSION} - safely organize documents "
                    f"by content.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "EXAMPLES\n"
            "  # Safe preview (default: dry-run, copy mode):\n"
            "  python smart_document_organizer.py \"C:\\Users\\me\\Downloads\"\n\n"
            "  # Preview with OCR for images and scanned PDFs, with renaming:\n"
            "  python smart_document_organizer.py \"C:\\Users\\me\\Downloads\" "
            "--use-ocr --scanned-pdf-ocr --rename\n\n"
            "  # Actually COPY files into an organized tree (originals kept):\n"
            "  python smart_document_organizer.py \"C:\\Users\\me\\Downloads\" "
            "--apply --copy --rename\n\n"
            "  # MOVE files instead of copying (originals relocated):\n"
            "  python smart_document_organizer.py \"C:\\Users\\me\\Downloads\" "
            "--apply --move\n\n"
            "  # Learn from your CSV corrections next time:\n"
            "  python smart_document_organizer.py \"C:\\Users\\me\\Downloads\" "
            "--learn-from \"Organized_Documents\\organizer_log.csv\" --apply\n\n"
            "  # Launch the GUI:\n"
            "  python smart_document_organizer.py --gui\n"
        ),
    )
    p.add_argument("folder", nargs="?", help="Source folder to scan "
                   "(recursively).")
    p.add_argument("--output", "-o", help="Output folder for the organized "
                   "tree. Default: <folder>/Organized_Documents.")
    p.add_argument("--apply", action="store_true",
                   help="Actually copy/move files. Without this it is a "
                        "dry-run (nothing is touched).")

    mode = p.add_mutually_exclusive_group()
    mode.add_argument("--copy", action="store_true",
                      help="Copy files (SAFE, recommended). This is the "
                           "default operation.")
    mode.add_argument("--move", action="store_true",
                      help="Move files instead of copying (originals are "
                           "relocated).")

    p.add_argument("--use-ocr", action="store_true",
                   help="Enable OCR for image files (JPG/PNG/TIFF/WEBP).")
    p.add_argument("--scanned-pdf-ocr", action="store_true",
                   help="Enable OCR for PDFs that contain no extractable text.")
    p.add_argument("--duplicate-action", choices=["keep", "folder", "skip"],
                   default="folder",
                   help="How to handle exact duplicates. keep=organize "
                        "normally; folder=put in _Duplicates (default); "
                        "skip=leave in place and only log.")
    p.add_argument("--min-confidence", type=float, default=0.45,
                   help="Minimum confidence (0-1) to file into a category; "
                        "below this goes to Needs Review. Default 0.45.")
    p.add_argument("--move-needs-review", action="store_true",
                   help="Also move/copy low-confidence files into "
                        "_Needs_Review. Without this they are only logged.")
    p.add_argument("--learn-from", metavar="CSV",
                   help="Path to a corrected CSV; apply corrected_category by "
                        "SHA-256 before classifying.")
    p.add_argument("--rename", action="store_true",
                   help="Rename files to "
                        "'YYYY-MM-DD - Category - Party - Hint.ext'.")
    p.add_argument("--ocr-lang", default="eng",
                   help="Tesseract language(s), e.g. 'eng' or 'eng+ara'.")
    p.add_argument("--reprocess", action="store_true",
                   help="Ignore the persistent history and process every "
                        "file again, even content organized by a previous "
                        "run (may create ' (2)' copies in the output).")
    p.add_argument("--review-by-type", action="store_true",
                   help="Group files inside _Needs_Review into subfolders "
                        "by file type, e.g. _Needs_Review\\PDF, "
                        "_Needs_Review\\MP3, _Needs_Review\\ZIP. "
                        "Use together with --move-needs-review.")
    p.add_argument("--find-duplicates", action="store_true",
                   help="Standalone mode: scan the folder (INCLUDING any "
                        "Organized_Documents inside it) for exact duplicate "
                        "files and write duplicates_report.csv. Report only "
                        "— nothing is touched.")
    p.add_argument("--delete-duplicates", action="store_true",
                   help="Like --find-duplicates, but also removes the "
                        "redundant copies after confirmation. One copy per "
                        "group is always kept (organized copies preferred). "
                        "Removed files go to the Recycle Bin (Send2Trash) "
                        "or a _Duplicates_Trash quarantine folder — never "
                        "permanently deleted.")
    p.add_argument("--yes", action="store_true",
                   help="Skip the confirmation prompt of "
                        "--delete-duplicates.")
    p.add_argument("--gui", action="store_true",
                   help="Launch the Tkinter graphical interface.")
    p.add_argument("--verbose", "-v", action="store_true",
                   help="Verbose logging.")
    p.add_argument("--version", action="version",
                   version=f"{APP_NAME} {APP_VERSION}")
    return p


def options_from_args(args: argparse.Namespace) -> Options:
    folder = Path(args.folder).expanduser()
    if args.output:
        output = Path(args.output).expanduser()
    else:
        output = folder / DEFAULT_OUTPUT_DIRNAME
    return Options(
        folder=folder,
        output=output,
        apply=args.apply,
        copy=not args.move,  # copy is default unless --move is given
        use_ocr=args.use_ocr,
        scanned_pdf_ocr=args.scanned_pdf_ocr,
        duplicate_action=args.duplicate_action,
        min_confidence=args.min_confidence,
        move_needs_review=args.move_needs_review,
        review_by_type=args.review_by_type,
        rename=args.rename,
        learn_from=Path(args.learn_from).expanduser() if args.learn_from
        else None,
        ocr_lang=args.ocr_lang,
        reprocess=args.reprocess,
    )


# --------------------------------------------------------------------------- #
# Standalone duplicate finder / cleaner (--find-duplicates / --delete-duplicates)
# --------------------------------------------------------------------------- #

DUP_REPORT_NAME = "duplicates_report.csv"
TRASH_DIRNAME = "_Duplicates_Trash"   # quarantine used when Send2Trash missing


def _human_size(n: float) -> str:
    for unit in ("B", "KB", "MB", "GB", "TB"):
        if n < 1024 or unit == "TB":
            return f"{n:,.0f} {unit}" if unit == "B" else f"{n:,.1f} {unit}"
        n /= 1024
    return f"{n:,.1f} TB"


def run_duplicate_scan(folder: Path, delete: bool = False,
                       assume_yes: bool = False) -> int:
    """Scan `folder` for exact duplicate files (SHA-256) and optionally
    remove the redundant copies.

    Differences from the organizer's normal duplicate handling:
      * EVERY file is scanned — all extensions, and the Organized_Documents
        tree is INCLUDED, so copies still sitting in the source can be
        matched against their already-organized twins.
      * One "keeper" per group is chosen and never touched. Preference:
        a copy inside Organized_Documents first, then the oldest file,
        then the shortest path.
      * With `delete=True`, redundant copies go to the RECYCLE BIN
        (Send2Trash). If Send2Trash isn't installed they are moved to a
        _Duplicates_Trash quarantine folder instead. Nothing is ever
        permanently deleted, so every removal is reversible.
    """
    # -- collect files ------------------------------------------------------ #
    files: List[Path] = []
    for root, dirs, fnames in os.walk(folder):
        root_path = Path(root)
        dirs[:] = [d for d in dirs
                   if d.lower() not in SKIP_DIR_NAMES and d != TRASH_DIRNAME]
        for fname in fnames:
            # Skip our own bookkeeping files — they change every run.
            if fname in (LOG_CSV_NAME, DB_NAME, DUP_REPORT_NAME):
                continue
            files.append(root_path / fname)

    logger.info("Hashing %d file(s) under %s ...", len(files), folder)
    by_hash: Dict[str, List[Path]] = {}
    errors = 0
    for i, p in enumerate(files, start=1):
        try:
            if p.stat().st_size == 0:
                continue  # all empty files are trivially "identical" — skip
            by_hash.setdefault(sha256_of_file(p), []).append(p)
        except Exception as exc:  # noqa: BLE001
            errors += 1
            logger.warning("Could not hash %s: %s", p, exc)
        if i % 200 == 0:
            logger.info("  ... %d / %d hashed", i, len(files))

    groups = {h: ps for h, ps in by_hash.items() if len(ps) > 1}

    # -- pick the keeper of each group -------------------------------------- #
    def keeper_rank(p: Path) -> Tuple[int, float, int]:
        in_organized = DEFAULT_OUTPUT_DIRNAME in p.parts
        try:
            mtime = p.stat().st_mtime
        except OSError:
            mtime = float("inf")
        return (0 if in_organized else 1, mtime, len(str(p)))

    dupes: List[Tuple[str, Path, Path]] = []  # (sha, duplicate, keeper)
    wasted = 0
    for sha, paths in groups.items():
        ordered = sorted(paths, key=keeper_rank)
        keeper = ordered[0]
        for extra in ordered[1:]:
            dupes.append((sha, extra, keeper))
            try:
                wasted += extra.stat().st_size
            except OSError:
                pass

    logger.info("Found %d duplicate group(s): %d redundant file(s), "
                "%s reclaimable.", len(groups), len(dupes),
                _human_size(wasted))

    # -- confirmation + removal --------------------------------------------- #
    removed_action = ""
    if delete and dupes:
        try:
            from send2trash import send2trash as _send2trash
        except ImportError:
            _send2trash = None
            logger.warning(
                "Send2Trash is not installed (pip install Send2Trash). "
                "Duplicates will be MOVED to the '%s' quarantine folder "
                "instead of the Recycle Bin.", TRASH_DIRNAME)

        if not assume_yes:
            print(f"\nAbout to remove {len(dupes)} duplicate file(s) "
                  f"({_human_size(wasted)}). One copy of each is kept. "
                  f"Removed files go to the "
                  f"{'Recycle Bin' if _send2trash else TRASH_DIRNAME + ' folder'}"
                  f" and can be restored.")
            answer = input("Type DELETE to confirm, anything else aborts: ")
            if answer.strip().upper() != "DELETE":
                logger.info("Aborted. Nothing was removed. The report will "
                            "still be written.")
                delete = False

    results: List[Tuple[str, Path, Path, str]] = []
    if delete and dupes:
        trash_root = folder / TRASH_DIRNAME
        for sha, extra, keeper in dupes:
            try:
                if _send2trash is not None:
                    _send2trash(str(extra))
                    action = "sent to Recycle Bin"
                else:
                    dest = unique_destination(trash_root, extra.name)
                    dest.parent.mkdir(parents=True, exist_ok=True)
                    shutil.move(str(extra), str(dest))
                    action = f"moved to quarantine: {dest}"
            except Exception as exc:  # noqa: BLE001
                action = f"ERROR: {exc}"
                logger.warning("Could not remove %s: %s", extra, exc)
            results.append((sha, extra, keeper, action))
        removed_action = "removed"
    else:
        results = [(sha, extra, keeper, "found (report only)")
                   for sha, extra, keeper in dupes]

    # -- report -------------------------------------------------------------- #
    report = folder / DUP_REPORT_NAME
    with open(report, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.writer(f)
        w.writerow(["duplicate_group", "sha256", "duplicate_file",
                    "size", "kept_original", "action"])
        for sha, extra, keeper, action in results:
            try:
                size = extra.stat().st_size
            except OSError:
                size = ""
            w.writerow([sha[:12], sha, str(extra), size, str(keeper), action])

    logger.info("Report written: %s", report)
    if errors:
        logger.info("%d file(s) could not be read (see warnings above).",
                    errors)
    if not delete and dupes:
        logger.info("Nothing was removed (report-only mode). To remove the "
                    "redundant copies, re-run with --delete-duplicates.")
    elif removed_action:
        logger.info("Done. Every removal is reversible (Recycle Bin or "
                    "quarantine folder).")
    return 0


def run_cli(args: argparse.Namespace) -> int:
    if not args.folder:
        logger.error("No source folder given. Pass a folder or use --gui. "
                     "Try -h for help.")
        return 2
    opt = options_from_args(args)
    if not opt.folder.exists() or not opt.folder.is_dir():
        logger.error("Source folder does not exist or is not a directory: %s",
                     opt.folder)
        return 2

    # Standalone duplicate-finder mode: no classification, no organizing.
    if args.find_duplicates or args.delete_duplicates:
        return run_duplicate_scan(opt.folder,
                                  delete=args.delete_duplicates,
                                  assume_yes=args.yes)

    # Loud, explicit banner about what mode we're in so nothing is a surprise.
    op = "COPY" if opt.copy else "MOVE"
    mode = "APPLY" if opt.apply else "DRY-RUN (no changes)"
    logger.info("%s v%s", APP_NAME, APP_VERSION)
    logger.info("Mode: %s | Operation: %s", mode, op)
    logger.info("Source: %s", opt.folder)
    logger.info("Output: %s", opt.output)
    if opt.apply and not opt.copy:
        logger.warning("MOVE mode will relocate your original files.")

    organizer = Organizer(opt)
    organizer.run()
    return 0


# --------------------------------------------------------------------------- #
#  Tkinter GUI (optional, thin wrapper over the same engine)
# --------------------------------------------------------------------------- #

def launch_gui() -> int:
    """A simple Tkinter GUI. The engine runs on a background thread so the UI
    stays responsive; progress is marshalled back via a thread-safe queue."""
    try:
        import tkinter as tk
        from tkinter import ttk, filedialog, messagebox
    except Exception as exc:  # noqa: BLE001
        logger.error("Tkinter is not available: %s", exc)
        return 1

    import queue
    import threading

    root = tk.Tk()
    root.title(f"{APP_NAME} v{APP_VERSION}")
    root.geometry("980x620")

    state = {"records": [], "output": None}
    ui_queue: "queue.Queue" = queue.Queue()

    # -- top controls ------------------------------------------------------- #
    frm = ttk.Frame(root, padding=10)
    frm.pack(fill="x")

    src_var = tk.StringVar()
    out_var = tk.StringVar()
    dup_var = tk.StringVar(value="folder")
    conf_var = tk.DoubleVar(value=0.45)
    ocr_var = tk.BooleanVar(value=False)
    spdf_var = tk.BooleanVar(value=False)
    rename_var = tk.BooleanVar(value=False)
    review_var = tk.BooleanVar(value=False)
    bytype_var = tk.BooleanVar(value=False)

    def pick_source():
        d = filedialog.askdirectory(title="Select folder to organize")
        if d:
            src_var.set(d)
            if not out_var.get():
                out_var.set(str(Path(d) / DEFAULT_OUTPUT_DIRNAME))

    def pick_output():
        d = filedialog.askdirectory(title="Select output folder")
        if d:
            out_var.set(d)

    ttk.Label(frm, text="Source folder:").grid(row=0, column=0, sticky="w")
    ttk.Entry(frm, textvariable=src_var, width=80).grid(row=0, column=1,
                                                        sticky="we", padx=5)
    ttk.Button(frm, text="Browse...", command=pick_source).grid(row=0,
                                                                column=2)

    ttk.Label(frm, text="Output folder:").grid(row=1, column=0, sticky="w")
    ttk.Entry(frm, textvariable=out_var, width=80).grid(row=1, column=1,
                                                        sticky="we", padx=5)
    ttk.Button(frm, text="Browse...", command=pick_output).grid(row=1,
                                                                column=2)

    opts = ttk.Frame(frm)
    opts.grid(row=2, column=0, columnspan=3, sticky="w", pady=(8, 0))
    ttk.Checkbutton(opts, text="OCR images", variable=ocr_var).pack(
        side="left")
    ttk.Checkbutton(opts, text="OCR scanned PDFs", variable=spdf_var).pack(
        side="left", padx=(10, 0))
    ttk.Checkbutton(opts, text="Rename files", variable=rename_var).pack(
        side="left", padx=(10, 0))
    ttk.Checkbutton(opts, text="Move Needs-Review too",
                    variable=review_var).pack(side="left", padx=(10, 0))
    ttk.Checkbutton(opts, text="Group review by type",
                    variable=bytype_var).pack(side="left", padx=(10, 0))
    ttk.Label(opts, text="Duplicates:").pack(side="left", padx=(10, 0))
    ttk.Combobox(opts, textvariable=dup_var, width=8, state="readonly",
                 values=["folder", "keep", "skip"]).pack(side="left")
    ttk.Label(opts, text="Min confidence:").pack(side="left", padx=(10, 0))
    ttk.Spinbox(opts, from_=0.0, to=1.0, increment=0.05, width=5,
                textvariable=conf_var).pack(side="left")

    frm.columnconfigure(1, weight=1)

    # -- progress + table --------------------------------------------------- #
    prog = ttk.Progressbar(root, mode="determinate")
    prog.pack(fill="x", padx=10, pady=(8, 0))

    status_var = tk.StringVar(value="Ready.")
    ttk.Label(root, textvariable=status_var).pack(anchor="w", padx=10)

    columns = ("name", "category", "confidence", "dup", "status")
    tree = ttk.Treeview(root, columns=columns, show="headings", height=18)
    for col, width, text in [
        ("name", 360, "File"),
        ("category", 180, "Category"),
        ("confidence", 90, "Confidence"),
        ("dup", 90, "Duplicate"),
        ("status", 220, "Status"),
    ]:
        tree.heading(col, text=text)
        tree.column(col, width=width, anchor="w")
    tree.pack(fill="both", expand=True, padx=10, pady=8)

    vsb = ttk.Scrollbar(tree, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=vsb.set)
    vsb.pack(side="right", fill="y")

    # -- worker thread ------------------------------------------------------ #
    def build_options(apply: bool, copy: bool) -> Optional[Options]:
        src = src_var.get().strip()
        if not src or not Path(src).is_dir():
            messagebox.showerror(APP_NAME, "Please choose a valid source "
                                           "folder.")
            return None
        out = out_var.get().strip() or str(Path(src) / DEFAULT_OUTPUT_DIRNAME)
        return Options(
            folder=Path(src),
            output=Path(out),
            apply=apply,
            copy=copy,
            use_ocr=ocr_var.get(),
            scanned_pdf_ocr=spdf_var.get(),
            duplicate_action=dup_var.get(),
            min_confidence=float(conf_var.get()),
            move_needs_review=review_var.get(),
            review_by_type=bytype_var.get(),
            rename=rename_var.get(),
        )

    def run_worker(opt: Options):
        def progress(i, total, rec: Record):
            ui_queue.put(("progress", (i, total, rec)))

        try:
            org = Organizer(opt, progress=progress)
            org.run()
            ui_queue.put(("done", (opt.output, org.records)))
        except Exception as exc:  # noqa: BLE001
            ui_queue.put(("error", str(exc)))

    def start(apply: bool, copy: bool):
        opt = build_options(apply, copy)
        if opt is None:
            return
        if apply and not copy:
            if not messagebox.askyesno(
                APP_NAME,
                "MOVE mode will relocate your original files. Continue?"
            ):
                return
        for item in tree.get_children():
            tree.delete(item)
        state["output"] = opt.output
        prog["value"] = 0
        status_var.set("Working...")
        threading.Thread(target=run_worker, args=(opt,), daemon=True).start()

    def poll_queue():
        try:
            while True:
                kind, payload = ui_queue.get_nowait()
                if kind == "progress":
                    i, total, rec = payload
                    prog["maximum"] = total
                    prog["value"] = i
                    status_var.set(f"Processing {i}/{total}: {rec.original_name}")
                    tree.insert("", "end", values=(
                        rec.original_name,
                        rec.category,
                        f"{rec.confidence:.2f}",
                        rec.duplicate_status,
                        rec.status,
                    ))
                elif kind == "done":
                    output, records = payload
                    state["records"] = records
                    status_var.set(
                        f"Done. {len(records)} file(s). Log: "
                        f"{Path(output) / LOG_CSV_NAME}"
                    )
                elif kind == "error":
                    status_var.set("Error.")
                    messagebox.showerror(APP_NAME, payload)
        except queue.Empty:
            pass
        root.after(100, poll_queue)

    # -- bottom buttons ----------------------------------------------------- #
    btns = ttk.Frame(root, padding=10)
    btns.pack(fill="x")
    ttk.Button(btns, text="Dry Run (preview)",
               command=lambda: start(apply=False, copy=True)).pack(side="left")
    ttk.Button(btns, text="Apply (Copy - safe)",
               command=lambda: start(apply=True, copy=True)).pack(
        side="left", padx=(8, 0))
    ttk.Button(btns, text="Apply (Move)",
               command=lambda: start(apply=True, copy=False)).pack(
        side="left", padx=(8, 0))

    def open_csv():
        if not state["output"]:
            return
        csv_path = Path(state["output"]) / LOG_CSV_NAME
        if csv_path.exists():
            _open_in_os(csv_path)
        else:
            messagebox.showinfo(APP_NAME, "No log yet. Run a scan first.")

    def open_output():
        if state["output"] and Path(state["output"]).exists():
            _open_in_os(Path(state["output"]))
        else:
            messagebox.showinfo(APP_NAME, "Output folder does not exist yet.")

    ttk.Button(btns, text="Open CSV log", command=open_csv).pack(
        side="right")
    ttk.Button(btns, text="Open output folder", command=open_output).pack(
        side="right", padx=(0, 8))

    poll_queue()
    root.mainloop()
    return 0


def _open_in_os(path: Path) -> None:
    """Open a file/folder with the OS default handler (cross-platform)."""
    try:
        if sys.platform.startswith("win"):
            os.startfile(str(path))  # type: ignore[attr-defined]
        elif sys.platform == "darwin":
            import subprocess
            subprocess.Popen(["open", str(path)])
        else:
            import subprocess
            subprocess.Popen(["xdg-open", str(path)])
    except Exception as exc:  # noqa: BLE001
        logger.error("Could not open %s: %s", path, exc)


# --------------------------------------------------------------------------- #
#  Entry point
# --------------------------------------------------------------------------- #

def main(argv: Optional[List[str]] = None) -> int:
    parser = build_arg_parser()
    args = parser.parse_args(argv)

    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(asctime)s  %(levelname)-7s  %(message)s",
        datefmt="%H:%M:%S",
    )

    if args.gui or (not args.folder and len(sys.argv) == 1):
        return launch_gui()
    return run_cli(args)


if __name__ == "__main__":
    raise SystemExit(main())
